"""
upload_routes.py — FastAPI router cho Admin upload tài liệu vào tri thức chatbox.

Endpoints (prefix /api/admin):
  POST /api/admin/upload    — upload file docx/pdf/txt/md → parse → chunk → ingest
  GET  /api/admin/check     — kiểm tra ADMIN_PASSWORD (Bearer)
  GET  /api/admin/knowledge — danh sách tài liệu đã upload (prefix upload/)
  POST /api/admin/delete    — xóa tài liệu theo file_path/title

Luồng upload:
  1) Xóa chunks cũ + insert knowledge_chunks (bảng chat production đọc)
  2) Upsert documents (bảng BM25 local đọc) → rebuild() để search thấy ngay
  3) Xóa toàn bộ answer_cache — tránh chatbox trả câu trả lời cũ
"""
import hmac
import logging
import os
import re
from io import BytesIO
from typing import Any

from fastapi import APIRouter, File, HTTPException, Request, UploadFile
from fastapi.responses import JSONResponse

from db import get_client, upsert_document
from search_engine import rebuild

logger = logging.getLogger("upload_routes")

router = APIRouter(prefix="/api/admin", tags=["admin-upload"])

ALLOWED_EXTENSIONS = {".docx", ".pdf", ".txt", ".md"}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB (local không bị giới hạn Vercel)
INSERT_BATCH = 50

# --- Regex (giống ingestion.py) ---
FRONTMATTER_RE = re.compile(r"^---\n(.*?)\n---\n?", re.DOTALL)
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
PARA_SPLIT_RE = re.compile(r"\n\s*\n")
MAX_CHUNK_TOKENS = 1500

# Ký tự không hợp lệ trong file_path → thay bằng '-'
SANITIZE_RE = re.compile(r"[\\/:*?\"<>|\s]+")

_ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "")


def _check_admin(request: Request) -> None:
    """Xác thực Bearer token so với ADMIN_PASSWORD (chống timing attack)."""
    if not _ADMIN_PASSWORD:
        raise HTTPException(status_code=401, detail="ADMIN_PASSWORD chưa được cấu hình")
    auth = request.headers.get("Authorization", "")
    if not auth.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Cần token quản trị")
    token = auth.replace("Bearer ", "", 1)
    if not hmac.compare_digest(token, _ADMIN_PASSWORD):
        raise HTTPException(status_code=401, detail="Token quản trị không hợp lệ")


def sanitize_title(title: str) -> str:
    """Làm sạch title: thay ký tự đặc biệt + dấu cách → '-', giới hạn 100 ký tự."""
    cleaned = SANITIZE_RE.sub("-", title).strip("-").replace("--", "-")
    cleaned = re.sub(r"-{2,}", "-", cleaned).strip("-")
    return cleaned[:100] or "untitled"


def count_tokens(text: str) -> int:
    return len(text.split())


def _chunk_by_heading(body: str) -> list[dict[str, str]]:
    """Cắt markdown body theo heading stack H1-H6 (port ingestion.py)."""
    lines = body.split("\n")
    sections: list[tuple[str, str]] = []
    heading_stack: list[str] = []
    current_lines: list[str] = []

    def flush() -> None:
        if current_lines:
            text = "\n".join(current_lines).strip()
            if text:
                sections.append((" > ".join(heading_stack), text))
            current_lines.clear()

    for line in lines:
        m = HEADING_RE.match(line)
        if m:
            flush()
            level = len(m.group(1))
            heading_stack = heading_stack[: level - 1]
            heading_stack.append(m.group(2).strip())
        else:
            current_lines.append(line)
    flush()

    chunks: list[dict[str, str]] = []
    for heading, text in sections:
        chunks.extend(_split_oversized(heading, text))
    return chunks


def _split_oversized(heading: str, text: str) -> list[dict[str, str]]:
    """Section quá MAX_CHUNK_TOKENS → cắt theo đoạn → theo từ."""
    if count_tokens(text) <= MAX_CHUNK_TOKENS:
        return [{"text": text, "heading": heading}]

    paragraphs = [p.strip() for p in PARA_SPLIT_RE.split(text) if p.strip()]
    chunks: list[dict[str, str]] = []
    buf: list[str] = []
    buf_tokens = 0

    for para in paragraphs:
        pt = count_tokens(para)
        if buf and buf_tokens + pt > MAX_CHUNK_TOKENS:
            chunks.append({"text": "\n\n".join(buf), "heading": heading})
            buf, buf_tokens = [], 0
        if pt > MAX_CHUNK_TOKENS:
            words = para.split()
            for i in range(0, len(words), MAX_CHUNK_TOKENS):
                chunks.append({"text": " ".join(words[i : i + MAX_CHUNK_TOKENS]), "heading": heading})
        else:
            buf.append(para)
            buf_tokens += pt

    if buf:
        chunks.append({"text": "\n\n".join(buf), "heading": heading})
    return chunks


def _chunk_plain_text(text: str) -> list[dict[str, str]]:
    """Fallback cho văn bản không heading: ≤1500 từ 1 chunk, quá thì cắt theo từ."""
    clean = text.strip()
    if not clean:
        return []
    words = clean.split()
    if len(words) <= MAX_CHUNK_TOKENS:
        return [{"text": clean, "heading": ""}]
    return [
        {"text": " ".join(words[i : i + MAX_CHUNK_TOKENS]), "heading": ""}
        for i in range(0, len(words), MAX_CHUNK_TOKENS)
    ]


def extract_text(filename: str, content: bytes) -> tuple[str, str, bool]:
    """Trích văn bản theo extension. Trả về (title mặc định, body, is_markdown)."""
    ext = os.path.splitext(filename)[1].lower()
    title = os.path.splitext(os.path.basename(filename))[0]

    if ext == ".docx":
        from docx import Document

        doc = Document(BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        # Gộp cả bảng nếu có (một số tài liệu thuế đặt nội dung trong table)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return title, "\n".join(parts), False

    if ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        parts: list[str] = []
        for page in reader.pages:
            try:
                text = page.extract_text() or ""
            except Exception:  # noqa: BLE001 — trang lỗi bỏ qua
                text = ""
            if text.strip():
                parts.append(text)
        return title, "\n\n".join(parts), False

    # .txt / .md
    return title, content.decode("utf-8", errors="replace"), ext == ".md"


def _clear_answer_cache(client: Any) -> None:
    """Xóa toàn bộ answer_cache — chatbox không trả câu trả lời cũ sau upload."""
    try:
        client.table("answer_cache").delete().neq(
            "id", "00000000-0000-0000-0000-000000000000"
        ).execute()
    except Exception as exc:  # noqa: BLE001 — best-effort
        logger.warning("Xóa answer_cache thất bại: %s", exc)


@router.post("/upload")
async def admin_upload(request: Request, file: UploadFile = File(...), title: str = ""):
    """Upload tài liệu → chunk → ingest knowledge_chunks + documents → rebuild BM25."""
    _check_admin(request)

    filename = file.filename or ""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Chỉ hỗ trợ .docx, .pdf, .txt, .md. File của bạn: {ext or 'không xác định'}",
        )

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File quá lớn ({len(content) / 1024 / 1024:.1f} MB). Tối đa 20 MB.",
        )

    try:
        default_title, body, is_markdown = extract_text(filename, content)
    except Exception as exc:  # noqa: BLE001
        logger.exception("Parse file thất bại: %s", filename)
        raise HTTPException(status_code=400, detail=f"Không đọc được nội dung file: {exc}") from exc

    if not body.strip():
        raise HTTPException(status_code=400, detail="File rỗng hoặc không đọc được nội dung")

    doc_title = (title or default_title).strip() or "untitled"
    file_path = "upload/" + sanitize_title(doc_title)

    # Chunk: markdown giữ heading, khác → plain text
    if is_markdown:
        fm = FRONTMATTER_RE.match(body)
        body_no_fm = body[fm.end():] if fm else body
        chunks = _chunk_by_heading(body_no_fm)
    else:
        chunks = _chunk_plain_text(body)

    if not chunks:
        raise HTTPException(status_code=400, detail="File rỗng hoặc không đọc được nội dung")

    client = get_client()

    # 1) Idempotent: insert (upsert theo file_path+chunk_index) TRƯỚC rồi xóa
    #    phần dư SAU (M4) — nếu insert fail (mạng/Supabase) thì chunks cũ vẫn
    #    còn, không mất file khỏi tri thức dù API báo 500.
    #    Nếu DB chưa có UNIQUE(file_path,chunk_index) (42P10) → fallback
    #    delete-cũ-trước + insert thuần để upload không gãy trên DB chưa migrate.
    plain_insert = False
    for i in range(0, len(chunks), INSERT_BATCH):
        batch = chunks[i : i + INSERT_BATCH]
        records = [
            {
                "content": c["text"],
                "title": doc_title,
                "heading": c["heading"],
                "file_path": file_path,
                "chunk_index": i + idx,
            }
            for idx, c in enumerate(batch)
        ]
        if plain_insert:
            client.table("knowledge_chunks").insert(records).execute()
            continue
        try:
            client.table("knowledge_chunks").upsert(
                records, on_conflict="file_path,chunk_index"
            ).execute()
        except Exception as exc:  # noqa: BLE001
            code = getattr(exc, "code", None)
            if code != "42P10" and "ON CONFLICT" not in str(exc):
                raise
            logger.warning(
                "DB thiếu UNIQUE(file_path,chunk_index) — fallback delete+insert (chạy migration 2026-08-01): %s",
                exc,
            )
            client.table("knowledge_chunks").delete().eq("file_path", file_path).execute()
            client.table("knowledge_chunks").insert(records).execute()
            plain_insert = True

    # Xóa các chunk dư (chunks mới ít hơn chunks cũ)
    client.table("knowledge_chunks").delete().eq("file_path", file_path).gte(
        "chunk_index", len(chunks)
    ).execute()

    # 2) Ghi documents (BM25 local đọc) + rebuild index
    try:
        upsert_document(file_path, doc_title, body, chunks)
        rebuild()
    except Exception as exc:  # noqa: BLE001
        logger.warning("Ghi documents/rebuild thất bại (chunks vẫn đã lưu): %s", exc)

    # 3) Xóa cache câu trả lời cũ
    _clear_answer_cache(client)

    logger.info("Upload xong: %s (%d chunks)", file_path, len(chunks))
    return JSONResponse(
        content={"ok": True, "chunks": len(chunks), "title": doc_title, "file_path": file_path}
    )


@router.get("/check")
async def admin_check(request: Request):
    """Kiểm tra ADMIN_PASSWORD hợp lệ."""
    _check_admin(request)
    return {"ok": True}


@router.get("/knowledge")
async def admin_knowledge(request: Request):
    """Danh sách tài liệu đã upload (prefix upload/) + số chunks."""
    _check_admin(request)
    client = get_client()
    res = (
        client.table("knowledge_chunks")
        .select("file_path,title,created_at")
        .like("file_path", "upload/%")
        .limit(2000)
        .execute()
    )
    rows = res.data or []

    docs: list[dict[str, Any]] = []
    seen: set[str] = set()
    for r in rows:
        fp = r.get("file_path", "")
        if fp in seen:
            continue
        seen.add(fp)
        cnt = (
            client.table("knowledge_chunks")
            .select("id", count="exact", head=True)
            .eq("file_path", fp)
            .execute()
        )
        docs.append(
            {
                "file_path": fp,
                "title": r.get("title") or fp,
                "chunk_count": cnt.count if cnt.count is not None else 0,
                "created_at": r.get("created_at") or "",
            }
        )

    docs.sort(key=lambda d: d["created_at"], reverse=True)
    return {"docs": docs}


@router.post("/delete")
async def admin_delete(request: Request):
    """Xóa tài liệu theo file_path (eq/ilike) hoặc title; đồng bộ xóa cache."""
    _check_admin(request)
    body = await request.json()
    file_path = body.get("file_path") or body.get("source") or ""
    title = body.get("title") or ""
    mode = body.get("mode", "contains")

    if not file_path and not title:
        raise HTTPException(status_code=400, detail="Thiếu file_path hoặc title")

    client = get_client()
    query = client.table("knowledge_chunks").delete(count="exact")
    if file_path:
        query = query.eq("file_path", file_path) if mode == "exact" else query.ilike(
            "file_path", f"%{file_path}%"
        )
    else:
        query = query.eq("title", title) if mode == "exact" else query.ilike(
            "title", f"%{title}%"
        )

    res = query.execute()
    deleted = res.count if res.count is not None else len(res.data or [])

    _clear_answer_cache(client)
    logger.info("Xóa xong: %s (%d chunks)", file_path or title, deleted)
    return {"ok": True, "deleted": deleted}
